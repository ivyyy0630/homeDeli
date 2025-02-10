from docx import Document
from datetime import datetime
from constants.path import OUTPUT_DOC_NAME, results_dir
import os

def clear_text(text):
    """
    去除字符串中的所有 </think> 标签。

    :param text: 输入的字符串
    :return: 去除 </think> 标签后的字符串
    """
    return text.replace("</think>", "").replace("<think>", "")

# 假设 response['text'] 是从 Ollama 模型获取的响应文本
def write_to_doc(response_text):

    # 检查 results 目录是否存在
    if not os.path.exists(results_dir):
        # 如果不存在，则创建目录
        os.makedirs(results_dir)
        print(f"目录 '{results_dir}' 已创建。")
    else:
        print(f"目录 '{results_dir}' 已存在。")
    # 创建一个新的 Word 文档
    doc = Document()

    # 添加标题
    doc.add_heading('Ollama 模型响应', level=1)

    # 添加段落
    doc.add_paragraph(clear_text(response_text))

    # 保存文档
    now = datetime.now()
    timestamp_str = now.strftime("%Y-%m-%d%H-%M-%S")
    output_path = OUTPUT_DOC_NAME + timestamp_str + '.docx'
    output_path = results_dir + '/' +  output_path
    doc.save(output_path)

    print(f"文档已保存到 {output_path}")
